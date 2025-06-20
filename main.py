import streamlit as st
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential
from azure.core.pipeline.policies import RetryPolicy
from openai import AzureOpenAI
from dotenv import load_dotenv
import re
import os
import base64
import json
from docx.shared import Pt
from bs4 import BeautifulSoup
import time
import hashlib
import uuid
from datetime import datetime
from rapidfuzz import process
from rapidfuzz import fuzz
import asyncio
from io import BytesIO
from docx import Document
from getdatafromblob import fetch_and_get_lesson,format_lesson_output
from dataformatting import convert_markdown_to_bold_html,convert_markdown_to_bold_html_1,convert_markdown_to_clean_text,convert_markdown_to_clean_text_for_docs
from log_to_blob import log_query_to_blob
from convert_to_pdf import generate_structured_pdf



def get_user_id():
    """Generate a simple user ID based on session properties"""
    if "user_id" not in st.session_state:
        st.session_state.user_id = str(uuid.uuid4())[:8]
    return st.session_state.user_id



async def async_azure_openai_call(messages):
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(
        None,
        lambda: client.chat.completions.create(
            model=OPENAI_DEPLOYMENT_NAME,
            messages=messages,
            temperature=0.9,
            max_tokens=16384
        )
    )


def initialize_session_history():
    """Initialize session history for the current user"""
    user_id = get_user_id()
    if "user_histories" not in st.session_state:
        st.session_state.user_histories = {}
    
    if user_id not in st.session_state.user_histories:
        st.session_state.user_histories[user_id] = []


def add_to_history(query, resource_id, benchmark, lesson_plan, ai_output):
    """Add a new entry to the user's history"""
    user_id = get_user_id()
    history_entry = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "query": query,
        "resource_id": resource_id,
        "benchmark": benchmark,
        "lesson_plan": lesson_plan,
        "ai_output": ai_output
    }
    
    print("=== ADDING TO HISTORY ===")
    print(f"User ID: {user_id}")
    print(f"Query: {query}")
    print(f"Resource ID: {resource_id}")
    print(f"Benchmark: {benchmark}")
    print(f"Lesson Plan Length: {len(lesson_plan) if lesson_plan else 0}")
    print(f"AI Output Length: {len(ai_output) if ai_output else 0}")
    print(f"Current History Length Before Add: {len(st.session_state.user_histories.get(user_id, []))}")
    
    st.session_state.user_histories[user_id].append(history_entry)
    st.session_state.user_histories[user_id] = st.session_state.user_histories[user_id][-10:]
    
    print(f"Current History Length After Add: {len(st.session_state.user_histories[user_id])}")
    print("=== END ADDING TO HISTORY ===")


def show_history():
    """Display the user's query history"""
    user_id = get_user_id()
    
    print("=== SESSION HISTORY DEBUG ===")
    print(f"Current User ID: {user_id}")
    print(f"User Histories Keys: {list(st.session_state.user_histories.keys()) if 'user_histories' in st.session_state else 'None'}")
    print(f"Current User History Length: {len(st.session_state.user_histories.get(user_id, []))}")
    
    if user_id in st.session_state.user_histories:
        print("User History Contents:")
        for i, entry in enumerate(st.session_state.user_histories[user_id]):
            print(f"  Entry {i+1}:")
            print(f"    Timestamp: {entry.get('timestamp', 'N/A')}")
            print(f"    Resource ID: {entry.get('resource_id', 'N/A')}")
            print(f"    Benchmark: {entry.get('benchmark', 'N/A')}")
            print(f"    Query Length: {len(entry.get('query', ''))}")
            print(f"    Lesson Plan Length: {len(entry.get('lesson_plan', ''))}")
            print(f"    AI Output Length: {len(entry.get('ai_output', ''))}")
    else:
        print("No history found for current user")
    print("=== END DEBUG ===")
    
    if st.session_state.user_histories[user_id]:
        st.markdown("### 📚 Previous Queries")
        for i, entry in enumerate(reversed(st.session_state.user_histories[user_id]), 1):
            with st.expander(f"Query {i}: {entry['timestamp']} - Resource ID: {entry['resource_id']}"):
                st.markdown("**Query:**")
                st.write(entry['query'])
                st.markdown("**Benchmark:**")
                st.write(entry['benchmark'])
                
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**📘 Original Lesson Plan**")
                    st.write(entry['lesson_plan'])
                with col2:
                    st.markdown("**✨ AI Customization**")
                    st.write(entry['ai_output'])


st.set_page_config(
    page_title="CPALMS AI Lesson Plan Generator",
    layout="wide",
    initial_sidebar_state="collapsed"
)

 
load_dotenv()
retry_policy = RetryPolicy(retry_total=2, timeout=120)
 
AZURE_SEARCH_ENDPOINT = os.getenv("AZURE_SEARCH_ENDPOINT")
AZURE_SEARCH_INDEX_NAME = os.getenv("AZURE_SEARCH_INDEX")
AZURE_SEARCH_INDEX_NAME_1 = os.getenv("AZURE_SEARCH_INDEX_1")
AZURE_SEARCH_API_KEY = os.getenv("AZURE_SEARCH_KEY")
 
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_API_VERSION = os.getenv("OPENAI_API_VERSION")
OPENAI_API_BASE = os.getenv("OPENAI_API_BASE")
OPENAI_DEPLOYMENT_NAME = os.getenv("OPENAI_DEPLOYMENT_NAME")
 
search_client = SearchClient(
    endpoint=AZURE_SEARCH_ENDPOINT,
    index_name=AZURE_SEARCH_INDEX_NAME,
    credential=AzureKeyCredential(AZURE_SEARCH_API_KEY),
    retry_policy=retry_policy
)
 
search_client_1 = SearchClient(
    endpoint=AZURE_SEARCH_ENDPOINT,
    index_name=AZURE_SEARCH_INDEX_NAME_1,
    credential=AzureKeyCredential(AZURE_SEARCH_API_KEY),
    retry_policy=retry_policy
)
 
client = AzureOpenAI(
    api_key=OPENAI_API_KEY,
    api_version=OPENAI_API_VERSION,
    azure_endpoint=OPENAI_API_BASE
)
 
allowed_benchmark_codes = {
    'ELA.1.R.1.4', 'ELA.1.R.3.1', 'ELA.4.V.1.3', 'ELA.5.V.1.1', 'ELA.5.V.1.3', 'ELA.6.V.1.3', 'ELA.7.C.1.3',
    'ELA.7.C.4.1', 'ELA.K.C.1.3', 'ELA.K.R.1.3', 'ELA.K.R.1.4', 'ELA.K.R.2.1', 'ELA.K.R.2.2', 'ELA.K.R.3.1',
    'ELA.K12.EE.1.1', 'ELA.K12.EE.2.1', 'ELA.K12.EE.3.1', 'ELA.K12.EE.4.1', 'ELA.K12.EE.6.1', 'MA.1.AR.1.1',
    'MA.1.GR.1.3', 'MA.1.NSO.1.1', 'MA.1.NSO.2.2', 'MA.1.NSO.2.4', 'MA.1.NSO.2.5', 'MA.2.AR.3.1', 'MA.2.AR.3.2',
    'MA.3.AR.1.1', 'MA.3.NSO.2.2', 'MA.3.NSO.2.4', 'MA.4.DP.1.2', 'MA.5.DP.1.2', 'MA.5.M.1.1', 'MA.5.NSO.2.4',
    'MA.5.NSO.2.5', 'MA.6.AR.3.2', 'MA.6.DP.1.2', 'MA.6.DP.1.3', 'MA.6.DP.1.4', 'MA.6.DP.1.5', 'MA.6.DP.1.6',
    'MA.6.GR.2.3', 'MA.6.GR.2.4', 'MA.6.NSO.2.3', 'MA.7.AR.3.1', 'MA.7.DP.1.1', 'MA.7.DP.1.2', 'MA.7.DP.1.5',
    'MA.7.DP.2.1', 'MA.8.F.1.3', 'MA.912.AR.1.3', 'MA.912.DP.1.1', 'MA.912.DP.1.2', 'MA.912.DP.1.4', 'MA.912.DP.2.1',
    'MA.912.DP.2.2', 'MA.912.DP.3.5', 'MA.912.T.3.3', 'MA.K.AR.1.1', 'MA.K.AR.1.2', 'MA.K.AR.1.3', 'MA.K.DP.1.1',
    'MA.K.GR.1.1', 'MA.K.GR.1.2', 'MA.K.GR.1.5', 'MA.K.M.1.2', 'MA.K.M.1.3', 'MA.K.NSO.1.1', 'MA.K.NSO.1.2',
    'MA.K.NSO.1.4', 'MA.K.NSO.2.1', 'MA.K.NSO.2.3', 'MA.K.NSO.3.1', 'MA.K.NSO.3.2', 'SS.7.CG.3.13', 'SS.7.CG.4.2',
    'SS.K.CG.2.2', 'SS.K.CG.2.4'
}
 


def convert_attachment_paths_to_links(paths):
    seen = set()
    unique_paths = []
    for path in paths:
        if path not in seen:
            seen.add(path)
            unique_paths.append(path)

    hyperlinks = []
    for i, url in enumerate(unique_paths, start=1):
        filename = os.path.basename(url)
        hyperlinks.append(f"{i}. [{filename}]({url})")
    return "\n".join(hyperlinks) 
def reset_session_state():
    """Reset session state for new query processing"""
    st.session_state.lesson_content = ""
    st.session_state.edit_mode = False
    st.session_state.copy_success = False
    st.session_state.show_copy_area = False


 
def has_query_changed():
    """Check if the current query is different from the last processed query"""
    current_query_key = f"{query}_{benchmark}_{resource_id}"
   
    if "last_query_key" not in st.session_state:
        st.session_state.last_query_key = ""
   
    if st.session_state.last_query_key != current_query_key:
        st.session_state.last_query_key = current_query_key
        return True
    return False
 
def should_process_query():
    """Determine if query should be processed (new query or empty lesson content)"""
    return not st.session_state.lesson_content or has_query_changed()
 
def clean_ai_response(text):
    """Clean AI response by removing separators, extra whitespace, and leading '#'"""
    text = re.sub(r'---\s*Chunk\s*\d+\s*Response\s*---', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    text = text.strip()
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = re.sub(r'^#+\s*', '', line)
        cleaned_line = line.strip()
        if cleaned_line and not re.match(r'^[-\s]*$', cleaned_line):
            cleaned_lines.append(line)
   
    return '\n'.join(cleaned_lines)
 

st.markdown("""
<style>
    /* Reset and base styles */
    .stApp {
        background-color: #f8f9fa !important;
    }
   
    /* Logo positioning */
    .logo-container {
        position: fixed;
        top: 20px;
        right: 20px;
        z-index: 1000;
        background: white;
        padding: 10px;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
    }
   
    /* Header styling */
    .main-header {
        text-align: center;
        padding: 40px 20px 20px 20px;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 15px;
        margin: 20px 0 30px 0;
        box-shadow: 0 10px 30px rgba(0,0,0,0.15);
    }
   
    .main-title {
        font-size: 2.8rem;
        font-weight: 700;
        margin-bottom: 10px;
        text-shadow: 0 2px 4px rgba(0,0,0,0.3);
    }
   
    .main-subtitle {
        font-size: 1.3rem;
        opacity: 0.9;
        font-weight: 300;
    }
   
    /* Input styling */
    .stTextInput > div > div > input {
        background-color: white;
        border: 2px solid #e1e5e9;
        border-radius: 10px;
        padding: 15px 20px;
        font-size: 16px;
        transition: all 0.3s ease;
        box-shadow: 0 2px 5px rgba(0,0,0,0.05);
    }
   
    .stTextInput > div > div > input:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
    }
   
    /* Action buttons */
    .action-buttons {
        display: flex;
        gap: 15px;
        margin: 25px 0;
        justify-content: center;
        flex-wrap: wrap;
    }
   
    /* Button styling improvements */
    .stButton > button {
        border-radius: 25px !important;
        font-weight: 600 !important;
        transition: all 0.2s ease !important;
        border: none !important;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1) !important;
    }
   
    .stButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 25px rgba(0,0,0,0.15) !important;
    }
   
    .stDownloadButton > button {
        background: linear-gradient(135deg, #FF9800, #F57C00) !important;
        color: white !important;
        border-radius: 25px !important;
        font-weight: 600 !important;
        transition: all 0.2s ease !important;
        border: none !important;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1) !important;
    }
   
    .stDownloadButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 25px rgba(0,0,0,0.15) !important;
    }
   
    /* Lesson content area */
    .lesson-display {
        background: white;
        border-radius: 15px;
        padding: 35px;
        margin: 25px 0;
        box-shadow: 0 5px 25px rgba(0,0,0,0.08);
        border: 1px solid #e6e6e6;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        line-height: 1.7;
        color: #333;
    }
   
    .lesson-display h1, .lesson-display h2, .lesson-display h3, .lesson-display h4 {
        color: #2c3e50;
        margin-top: 30px;
        margin-bottom: 15px;
        font-weight: 600;
    }
   
    .lesson-display h1 {
        font-size: 2rem;
        border-bottom: 3px solid #667eea;
        padding-bottom: 10px;
    }
   
    .lesson-display h2 {
        font-size: 1.5rem;
        color: #667eea;
    }
   
    .lesson-display strong, .lesson-display b {
        color: #2c3e50;
        font-weight: 600;
    }
   
    .lesson-display a {
        color: #667eea;
        text-decoration: none;
        font-weight: 500;
        transition: color 0.3s ease;
    }
   
    .lesson-display a:hover {
        color: #5a67d8;
        text-decoration: underline;
    }
   
    /* Edit mode styling */
    .edit-mode {
        border: 3px dashed #2196F3;
        background: #f8f9ff;
    }
   
    .stTextArea > div > div > textarea {
        border-radius: 10px;
        border: 2px solid #e1e5e9;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        font-size: 14px;
        line-height: 1.6;
    }
   
    .stTextArea > div > div > textarea:focus {
        border-color: #2196F3;
        box-shadow: 0 0 0 3px rgba(33, 150, 243, 0.1);
    }
   
    /* Warning messages */
    .warning-msg {
        background: linear-gradient(135deg, #ff6b6b, #ee5a52);
        color: white;
        padding: 20px 25px;
        border-radius: 12px;
        margin: 25px 0;
        box-shadow: 0 5px 20px rgba(238, 90, 82, 0.3);
        font-weight: 500;
    }
   
    /* Success messages */
    .success-msg {
        background: linear-gradient(135deg, #4CAF50, #45a049);
        color: white;
        padding: 15px 20px;
        border-radius: 10px;
        margin: 15px 0;
        box-shadow: 0 3px 15px rgba(76, 175, 80, 0.3);
        font-weight: 500;
        text-align: center;
    }
    
    /* Hide default streamlit elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .stDeployButton {display: none;}
   
    /* Loading spinner improvements */
    .stSpinner > div {
        border-color: #667eea transparent #667eea transparent !important;
    }
   
    /* Responsive design */
    @media (max-width: 768px) {
        .main-title {
            font-size: 2.2rem;
        }
       
        .action-buttons {
            justify-content: center;
        }
       
        .lesson-display {
            padding: 25px 20px;
        }
       
        .logo-container {
            position: relative;
            top: auto;
            right: auto;
            margin: 10px auto;
            text-align: center;
        }
    }
</style>
""", unsafe_allow_html=True)
 

if os.path.exists("/home/rohan/Downloads/CPalms 2/CPalms/cpalmslogo.png"):
    st.markdown("""
    <div class="logo-container">
        <img src="data:image/png;base64,{}" width="80">
    </div>
    """.format(base64.b64encode(open("/home/rohan/Downloads/CPalms 2/CPalms/cpalmslogo.png", "rb").read()).decode()),
    unsafe_allow_html=True)
 



def fuzzy_match_any_word(query, keywords, threshold):
    words = re.findall(r'\w+', query.lower())
    for word in words:
        match, score, _ = process.extractOne(word, keywords, scorer=fuzz.ratio)
        if score >= threshold:
            return True
    return False

 
def validate_educational_query(query: str) -> tuple[bool, str]:
    """
    Validate if the query is education-related and appropriate for lesson planning.
    Returns (is_valid, error_message)
    """
    query_lower = query.lower()
    
    educational_keywords = [
        'lesson', 'teaching', 'learning', 'student', 'classroom', 'activity', 'assessment', 
        'question', 'instruction', 'practice', 'exercise', 'worksheet', 'curriculum',
        'education', 'academic', 'school', 'grade', 'objective', 'skill', 'concept',
        'homework', 'assignment', 'project', 'discussion', 'explanation', 'example',
        'strategy', 'method', 'approach', 'technique', 'guidance', 'support', 'help',
        'understand', 'learn', 'study', 'review', 'prepare', 'develop', 'improve',
        'compare','help','prior knowledge','exam','plan','phases','collaborative activities',
        'knowledge', 'comprehension', 'mastery', 'benchmark', 'standard', 'goal',
        'outcome', 'performance', 'progress', 'achievement', 'rubric', 'criteria','quiz',
        'engage', 'explore', 'explain', 'elaborate', 'evaluate',
        'lesson plan', 'activity sheet', 'outcomes',
        'formative', 'summative', 'differentiation', 'scaffold', 'modification',
        'tactile', 'visual', 'auditory', 'kinesthetic','activity',
        'station', 'task', 'modeling', 'demonstration',
        'group work', 'pair work', 'independent work','learning stations'
    ]
    
    inappropriate_keywords = [
        'celebrity', 'gossip', 'politics', 'religion', 'personal',
        'dating', 'financial advice', 'medical advice', 'legal advice','gun','weapon',
        'inappropriate', 'violence', 'drugs', 'alcohol', 'gambling', 'adult content',
        'stock','investment'
    ]
    
        
    if fuzzy_match_any_word(query_lower, inappropriate_keywords, threshold=96):
        return False, "❌ This query contains inappropriate or off-topic content. Please focus on educational content such as lesson plans, activities, or assessments."


    if not fuzzy_match_any_word(query_lower, educational_keywords, threshold=70):
        return False, "❌ This query doesn't appear to be education-related. Please ask about lesson plans, teaching strategies, assessments, activities, or other educational content."

    
    return True, ""



def extract_required_section_from_query(query: str) -> list:
    """
    Extracts section types from a user's natural language query, such as
    assessments, activities, stations, prior knowledge, etc.
    """
    keywords = {
        "assessments": [
            "assessment", "assessments", "assessment questions", "quiz", "quizzes",
            "formative assessment", "summative assessment", "assessment games", "assessment rubrics"
        ],
        "prior_knowledge": [
            "prior knowledge", "prior knowledge requirements", "prior knowledge checklist","plan"
        ],
        "stations": [
            "station", "stations", "learning stations", "study stations", "problem-solving stations",
            "collaborative stations", "peer review stations", "station rotations"
        ],
        "activities": [
            "activity", "activities", "hands-on activities", "interactive activities",
            "collaborative activities", "creative workshops", "movement-based learning",
            "real-world applications", "context activities", "skill-building games", "investigations"
        ],
        "guiding_questions": [
            "guiding questions", "guiding question"
        ]
    }
 
    matched_sections = []
    q_lower = query.lower()
 
    for section, trigger_phrases in keywords.items():
        for phrase in trigger_phrases:
            if phrase in q_lower:
                matched_sections.append(section)
                break  
 
    return matched_sections
 
 




def initialize_session_state():
    """Safely initialize all session state variables used throughout the app"""
    defaults = {
        "lesson_content": "",
        "edit_mode": False,
        "copy_success": False,
        "last_processed_query": "",
        "show_copy_area": False,
        "user_histories": {},
        "user_id": str(uuid.uuid4())[:8],
        "lesson_plan_output": "",
        "last_query_key": ""
    }
    for key, default in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default


def generate_docx_file(content: str, title: str = "CPALMS Lesson Plan"):
    content = re.sub(r'\n\s*\n+', '\n', content.strip())

    doc = Document()
    doc.add_heading(title, level=0)

    for para in content.split("\n"):
        if not para.strip():
            doc.add_paragraph("")
            continue

        paragraph = doc.add_paragraph()
        while "**" in para:
            before, rest = para.split("**", 1)
            bold_text, after = rest.split("**", 1)
            paragraph.add_run(before)
            run = paragraph.add_run(bold_text)
            run.bold = True
            para = after
        paragraph.add_run(para)  

    return doc

def extract_test_or_worksheet_section(text: str) -> str:
    """
    Extract the section of the AI output that includes a worksheet, quiz, or test.
    Looks for headings like '## Worksheet' or '## Quiz Questions' and captures everything
    until the next heading or end of text.
    """
    pattern = r"(##\s*(Worksheet|Quiz|Test)[\s\S]*?)(?=\n##|\Z)"
    match = re.search(pattern, text, flags=re.IGNORECASE)
    
    if match:
        return match.group(1).strip()
    else:
        question_lines = []
        for line in text.splitlines():
            if any(q in line.lower() for q in ["question", "?", "1.", "a)", "b)", "answer"]):
                question_lines.append(line)
        return "\n".join(question_lines).strip()
    

def remove_inline_download_links(text: str) -> str:
    return re.sub(
        r'📄.*?\(data:application\/vnd\.openxmlformats-officedocument\.wordprocessingml\.document;base64,[^)]+\)',
        '', 
        text
    )

def make_docx_link(doc_buffer):
    doc_buffer.seek(0)
    b64 = base64.b64encode(doc_buffer.read()).decode()
    return f'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}'
def replace_generate_docx_link(markdown_text, doc_buffer):
    data_uri = make_docx_link(doc_buffer)
    return re.sub(r'\[(.*?)\]\(#GENERATE_DOCX_LINK\)', rf'[\1]({data_uri})', markdown_text)

    
def create_query_form():
    with st.form(key="query_form", clear_on_submit=False):
        col1, col2, col3 = st.columns([1, 1, 1])
        with col1:
            resource_id_input = st.text_input("🔢 Resource ID", value="26646", placeholder="e.g. 26646", key="resource_id_input")
        with col2:
            benchmark_code_input = st.text_input("🧠 Benchmark Code", value="MA.K.NSO.1.1", placeholder="e.g. MA.K.NSO.1.1", key="benchmark_code_input")
        with col3:
            benchmark_id_input = st.text_input("🆔 Benchmark ID", value="15232", placeholder="e.g. 15232", key="benchmark_id_input")

        query_input = st.text_input(
            "📝 Detailed Request",
            value="Generate teaching phase and guiding questions for students struggling with counting",
            placeholder="Example: Generate teaching phase and guiding questions for students struggling with counting...",
            key="query_input"
        )

        col_submit = st.columns([3, 2, 3])[1]
        with col_submit:
            submit_button = st.form_submit_button("🚀 Generate Lesson Plan", use_container_width=True,type="primary")

        return resource_id_input, benchmark_code_input, benchmark_id_input, query_input, submit_button



def normalize_benchmark_code(user_input):
    cleaned = re.sub(r'[^A-Za-z0-9]', '', user_input).upper()
    if cleaned in [re.sub(r'[^A-Za-z0-9]', '', b) for b in allowed_benchmark_codes]:
        for code in allowed_benchmark_codes:
            if cleaned == re.sub(r'[^A-Za-z0-9]', '', code):
                return code

    cleaned_allowed = {re.sub(r'[^A-Za-z0-9]', '', b): b for b in allowed_benchmark_codes}
    match, score, _ = process.extractOne(
        cleaned,
        list(cleaned_allowed.keys()),
        scorer=fuzz.ratio
    )

    if score >= 80:
        return cleaned_allowed[match]

    return None



def should_process_new_query(query, resource_id, benchmark_code, benchmark_id):
    current_query_key = f"{query}_{resource_id}_{benchmark_code}_{benchmark_id}"

    if "last_query_key" not in st.session_state:
        st.session_state.last_query_key = ""

    if st.session_state.last_query_key != current_query_key:
        st.session_state.last_query_key = current_query_key
        reset_session_state()
        return True
    return False
 
 
def generate_creative_response(query, des, grade_level, resource_id, benchmark, attachments_hyperlinks, docs_data, combined_chunks):
    """
    Generate a creative, comprehensive response for the specific question asked,
    using available lesson plan context and search results.
    """
    user_id = get_user_id()
    user_history = st.session_state.user_histories.get(user_id, [])
    
    history_context = ""
    previous_full_response = ""
    is_follow_up = False
    
    if user_history:
        current_session = [entry for entry in user_history 
                          if entry['resource_id'] == resource_id and entry['benchmark'] == benchmark]
        if current_session:
            is_follow_up = True
            latest_entry = current_session[-1]
            previous_entry = latest_entry

            history_context = ""

            if previous_entry:
                prev_response = previous_entry["ai_output"]

                split_parts = prev_response.split("📘 **Previous Response**")

                if len(split_parts) > 1:
                    new_content_only = split_parts[0].strip()
                    history_context += f"**Previous Response (New Content Portion Only):**\n{new_content_only}\n\n"
                else:
                    history_context += f"**Previous Response:**\n{prev_response.strip()}\n\n"

        else:
            history_context = "\n**Previous Session Context:**\n"
            for i, entry in enumerate(user_history[-2:], 1):  # Include last 2 queries for context
                history_context += f"{i}. Previous Query: {entry['query']}\n"
                history_context += f"   Resource ID: {entry['resource_id']}, Benchmark: {entry['benchmark']}\n"
                history_context += f"   Previous Response Summary: {entry['ai_output'][:300]}...\n\n"
    history_context=remove_inline_download_links(history_context)
    print("=== HISTORY CONTEXT SENT TO OPENAI ===")
    print(f"User ID: {user_id}")
    print(f"Number of history entries: {len(user_history)}")
    print("History context being sent:")
    print(history_context if history_context else "No history context")
    print("=== END HISTORY CONTEXT ===")
    
    system_content = f"""
You are a creative educational content generator specializing in lesson plan enhancement for CPALMS (Collaborative Planning for Learning in Mathematics and Science).

**STRICT OPERATIONAL GUIDELINES:**
- You MUST ONLY respond to queries related to education, lesson planning, teaching strategies, assessments, questions,quiz and classroom activities
- You MUST NOT respond to queries about weather reports, sports reports, celebrities, politics, personal advice, medical/legal advice, or any non-educational topics
- If a query is not education-related respond with: "I can only assist with educational content and lesson planning. Please ask about teaching strategies, assessments, activities,quiz or other lesson plan components."
- ALL responses must be directly related to the provided Resource ID, Benchmark, {query},{des} and {grade_level}.
- Stay on-topic and within educational context at all times and never attempt to interpret non-educational requests as educational.
- Generate the data based on only the {des} and {grade_level} only.example:for "grade :K,des:students knows numbers upto 5,response:should contain numbers till 5 not beyond."
- If user asks to create a worksheet, quiz, or test (e.g., "create 10-question test"), you MUST generate the actual content (not just suggestions), formatted with questions and answer choices where appropriate.
- If the user requests or if the content logically involves a worksheet, quiz, or test (even implicitly), DO NOT suggest that the teacher create one. INSTEAD, generate the full worksheet/test content directly.
- Never use or echo {attachments_hyperlinks}. Instead, insert this exact placeholder for downloads:[📄 Download Worksheet as doc](#GENERATE_DOCX_LINK).Always include it at the end or with the content.

**Context:**
- Resource ID: {resource_id}
- Benchmark: {benchmark}
- Available lesson plan sections: Learning Objectives, Prior Knowledge, Guiding Questions, Teaching Phase, Guided Practice, Independent Practice, Closure, Assessments, Accommodations, etc.
- Attachments: {attachments_hyperlinks}
- Follow-up Request: {"YES - This is a follow-up request" if is_follow_up else "NO - This is a new request"}

{history_context}

**Your Task:**
Analyze the user's specific request: "{query}"
Analyze the following lesson description, grade level and {query}, then generate creative content aligned with it:\n\ngrade Level: {grade_level}\ndescription: {des}

**CRITICAL INSTRUCTIONS BASED ON REQUEST TYPE:**

{"**THIS IS A FOLLOW-UP REQUEST:**" if is_follow_up else "**THIS IS A NEW REQUEST:**"}

{'''
- You MUST include the *COMPLETE* previous response in your output.
- HOWEVER, display the NEW content FIRST, and the previous response BELOW it satisfying -**If the user's query asks to *remove* or *exclude* previous responses (e.g., "ignore previous", "remove old content", "start fresh"), then do NOT include the ## 📘 **Previous Response** or if ## 📘 **Previous Response** has multiple sections remove that mentioned section and continue.**
- **If the user's query asks to *remove* or *exclude* previous responses (e.g., "ignore previous", "remove old content", "start fresh"), then do NOT include the ## 📘 **Previous Response** or remove that particular section in previous response and continue. Just generate new content as if this were a new request.**
- If user specifies to "remove" specific section then remove it from previous response. 
 - Use clear headers and spacing to visually separate them.
- For example if previous response is there then,
    ## ✨ **Latest Customization**
    [new additions]
        
 
    ## 📘 **Previous Response**
    [complete previous content]
- This layout ensures the newest information is always shown at the top of the output it is mandatory to use the same format if previous response is present.
- This ensures the user sees both the original content AND the new additions in one complete response
- Maintain consistency with the previous response's style and format
- DO NOT reuse wording or content from the attachments or provided document context unless explicitly asked to summarize it.
- For worksheets, quizzes, or tests, you MUST generate **original content** that is not found in the attachments or lesson data.
- Do not include file names or references from the attachments unless the user says “use that file.”
- If asked for more of the same type (e.g., "add more stations"), continue numbering from where the previous response ended''' if is_follow_up else '''- Generate complete, comprehensive content from scratch
- Provide thorough coverage of the requested topic
- Create standalone content that doesn't assume previous context
- Be comprehensive and detailed in your initial response'''}



**Content Guidelines:**
1. **Be creative and engaging** - Use varied teaching strategies, real-world connections, and student-centered approaches  
2. **Use provided data as foundation** - Build upon existing lesson content when relevant
3. **Format appropriately** - Use clear headings, bullet points, and structured content
4. **Be comprehensive** - Provide detailed, actionable content (minimum 1000 words per section requested)
5. **Include practical examples** - Give specific activities, questions, or scenarios
6. ****Strict grade-level appropriateness** - Ensure all content is developmentally appropriate for the specified grade level. For Kindergarten, avoid complex word problems, multi-step logic, or real-world contexts that require abstract thinking. Use simple language, visual elements, and tactile-friendly examples.**


**For specific request types:**
- **Assessments**: Create varied question types (multiple choice, short answer, performance tasks) - **minimum 10 questions**. DO NOT reuse questions from attachments. Provide original questions. At the end, insert a link like: [📄 Download Worksheet as doc](#GENERATE_DOCX_LINK)
- **Activities**: Design hands-on, collaborative, and differentiated activities
- **Stations**: Create 3-5 distinct learning stations with clear objectives and descriptions  
- **Prior Knowledge**: Identify prerequisites and diagnostic strategies (minimum 2000 words if requested)
- **Guiding Questions**: Develop thought-provoking, inquiry-based questions
- Do NOT use any attachments_hyperlinks. Instead, insert this link:[📄 Download Worksheet as doc](#GENERATE_DOCX_LINK)


### 📄 Worksheet / Assessment DOC Generation Rules
- Never say “create a worksheet” or suggest that a teacher prepares one. Always generate it directly and strictly based on {grade_level} and {des} instructions given above.
- If needed, say exactly:
    You can use the following worksheet with students:
    [📄 Download Worksheet as doc](#GENERATE_DOCX_LINK)
- Then generate the full content in the DOC. Adjust the format based on the type:
  - **Worksheets**: Start with **Name:** ________  **Date:** ________, then 10+ grade-appropriate questions, end with an **Answer Key**
  - **Quizzes**: Include clear questions (MCQs, short answer), instructions, and an Answer Key
  - **Plans**: Use a structured format with headings, instructions, and space for student responses
- Generate the DOC if:
  1. The user requests a worksheet, quiz, or assessment “as doc”
  2. Your response recommends using one (e.g., “students should complete a worksheet”)
- Everything should be matched according to grade only. like if it is for Kindergarten,they cannot solve word questions.


### Additional Enforcement Rules:
- For Kindergarten:
  - **Focus on visual, tactile, symbolic, or object-based interactions only — such as sorting, counting pictures, or matching icons (e.g., 🍎 + 🍌 = ?)**.
  - You must only generate questions and content based on the provided lesson {des} and {grade_level}. Do NOT introduce standards, objectives, or math skills beyond what is described.
  - Do NOT include word problems, narrative questions, or reading-based scenarios unless the description clearly states that students are ready for them.
  - Do NOT include number combinations or equations involving values beyond the specified range. For example, if students are learning to count or sort objects, do not include making 10, subtraction, or addition beyond 5 unless described.
- If a worksheet, quiz, or assessment is generated as a downloadable document, its content must NOT be repeated in the AI Customization response area. The AI customization must contain supporting or instructional content only — not the exact worksheet/quiz content — unless the user explicitly requests the questions be shown in both places.
- If the user's query requests to “remove previous,” “exclude old content,” “start fresh,” or similar — you MUST fully omit the `## 📘 Previous Response` block and all of its contents, including any quizzes, assessments, stations, or other structured sections.
- If the user specifies to “remove” or “replace” only a **specific section** (e.g., “remove previous quiz only” or “remove station 3”), then remove only that part from the `## 📘 Previous Response` block and regenerate it as new inside the `## ✨ Latest Customization` section. Keep the rest of the previous response unchanged.


**Output Format:**
Provide clean, well-organized content with clear section headers and practical details.
Use markdown formatting for better readability.
"""

    messages = [
        {"role": "system", "content": system_content},
        {"role": "user", "content": f"""
REMINDER: Only respond to educational queries related to lesson planning, teaching, assessments, or classroom activities. Refuse any non-educational requests.

User Query: {query}

Available Lesson Data: {docs_data}

Additional Context from Attachments: {combined_chunks}

{'IMPORTANT: This is a follow-up request. Your response can include the COMPLETE previous response shown in the system context above untill it is explicitly mentioned "remove" in {query}, followed by the new content. Start with the full previous content, then add a separator, then the new additions.' if is_follow_up else 'Generate creative, comprehensive content specifically for:'} {query}
"""}
    ]
    
    return messages

start_time = time.time()
 
st.markdown("""
<div class="main-header">
    <div class="main-title">🎓 CPALMS AI Lesson Plan Customizer</div>
    <div class="main-subtitle">Enhance lesson plan with AI assistance</div>
</div>
""", unsafe_allow_html=True)
 

initialize_session_state()
initialize_session_history()
 
st.markdown("### 💬 Enter Your Query")
resource_id_input, benchmark_code_input, benchmark_id_input, query, submit_clicked = create_query_form()

if submit_clicked:
    resource_id_input = resource_id_input.strip()
    benchmark_code_input = benchmark_code_input.strip()
    benchmark_id_input = benchmark_id_input.strip()
    query = query.strip()

    if not all([resource_id_input, benchmark_code_input, benchmark_id_input, query]):
        st.error("⚠️ All fields are required. Please fill out Resource ID, Benchmark Code, Benchmark ID, and Query.")
        st.stop()

    if not re.fullmatch(r'\d{5,6}', resource_id_input):
        st.error("❌ Resource ID must be a 5- or 6-digit number (only digits allowed).")
        st.stop()
    
    benchmark = normalize_benchmark_code(benchmark_code_input)
    if not benchmark:
        st.error("❌ Please enter a valid benchmark.")
        st.stop()
    
    is_valid_query, error_message = validate_educational_query(query)
    if not is_valid_query:
        st.error(error_message)
        st.stop()

if not submit_clicked and not st.session_state.lesson_content:
    st.stop()
 
query = query.strip()
 
if not query and st.session_state.lesson_content:
    benchmark = ""
    resource_id = ""
else:
    requested_sections = extract_required_section_from_query(query)
   
    if not query:
        st.stop()
 
query_upper = query.upper()
query_words = query_upper.split()
benchmark1=benchmark_code_input.strip()
benchmark = ""
benchmark=normalize_benchmark_code(benchmark1)
resource_id=resource_id_input.strip()
print(resource_id)
print(benchmark) 
matched_benchmarks = [word for word in query_words if word in allowed_benchmark_codes]
 

 
if resource_id is None:
    st.markdown("""
    <div class="warning-msg">
        <strong>📋 Resource ID Required</strong><br>
        Please provide a specific Resource ID (5-6 digits).
    </div>
    """, unsafe_allow_html=True)
    st.stop()
 
lesson_output_1 = fetch_and_get_lesson(benchmark, resource_id)
if isinstance(lesson_output_1, str) and lesson_output_1.startswith("⚠️"):
    st.warning(lesson_output_1) 
    st.stop() 
des=lesson_output_1.get("Description")
grade_level=lesson_output_1.get("GradeLevelNames")
title=lesson_output_1.get("Title")
objectives_resource_ids = set()
cnt=0
if submit_clicked and (should_process_new_query(query, resource_id, benchmark_code_input, benchmark_id_input) or not st.session_state.lesson_content):

    
    with st.spinner('🔄 Processing your request...'):
        query_for_resource = f"{resource_id} give all documents for this id"
        cnt = 0
       
        matched_docs = []
        search_results = search_client.search(search_text=benchmark, top=60)
        search_results = list(search_results)
 
        for i, doc in enumerate(search_results):
            doc_benchmarks = doc.get("benchmarkId", "")
            if benchmark in doc_benchmarks:
                doc_str = str(doc).lower()
                if any(section in doc_str for section in requested_sections):
                    filtered_doc = {k: v for k, v in doc.items() if k == "objectives"}
                    matched_docs.append(filtered_doc)
 
        attachments = []
        combined_chunks = ""
        search_results_1 = search_client_1.search(search_text=query_for_resource, top=60)
        for doc in search_results_1:
            path = doc.get("metadata_storage_path", "")
            match = re.search(r"/(\d{5,6})/", path)
            if match and match.group(1) == resource_id:
                attachments.append(path)
                chunk = doc.get("chunk", "")
                cnt += 1
                combined_chunks += chunk + "\n\n"
 
        attachments_hyperlinks = convert_attachment_paths_to_links(attachments)
        attachments_hyperlinks_list = attachments_hyperlinks.split("\n") if attachments_hyperlinks else []
 
        if cnt > 0:
            st.markdown(f"**📁 Retrieved data from {cnt} attachment(s):**")
            for link in attachments_hyperlinks_list:
                if link.strip():  # Only show non-empty links
                    st.markdown(f"- {link}", unsafe_allow_html=True)
        else:
            st.warning("⚠️ No attachments found for this Resource ID")
 
 
        
        docs_text = "\n\n".join([str(doc) for doc in matched_docs])
        q1=query+" targeted at Grade: "+grade_level
        q2=q1+" having title:"+title
        messages = generate_creative_response(
            query=q2,
            des=des,
            grade_level=grade_level,
            resource_id=resource_id,
            benchmark=benchmark,
            attachments_hyperlinks=attachments_hyperlinks,
            docs_data=docs_text,
            combined_chunks=combined_chunks
        )
        
        
        response = asyncio.run(async_azure_openai_call(messages))
        formatted_lesson = format_lesson_output(lesson_output_1,attachments_hyperlinks)
        st.session_state.lesson_plan_output = formatted_lesson

        lesson_output = response.choices[0].message.content
        if "#GENERATE_DOCX_LINK" in lesson_output:
            worksheet_section = extract_test_or_worksheet_section(lesson_output)
            worksheet_clean = convert_markdown_to_clean_text_for_docs(worksheet_section)
            doc = generate_docx_file(worksheet_clean, title="Student Worksheet")
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            lesson_output = replace_generate_docx_link(lesson_output, doc_io)
            st.session_state["worksheet_docx"] = doc_io


        lesson_output = clean_ai_response(lesson_output)
        st.session_state.lesson_content = lesson_output

        
        add_to_history(
            query=query,
            resource_id=resource_id,
            benchmark=benchmark,
            lesson_plan=st.session_state.lesson_plan_output,
            ai_output=st.session_state.lesson_content
        )
                

if st.session_state.lesson_content:
    formatted_lesson = convert_markdown_to_clean_text(st.session_state.lesson_plan_output)
    formatted_ai = remove_inline_download_links(convert_markdown_to_clean_text(st.session_state.lesson_content))
    formatted_lesson_for_docs = convert_markdown_to_clean_text_for_docs(st.session_state.lesson_plan_output)
    formatted_ai_for_docs = remove_inline_download_links(convert_markdown_to_clean_text_for_docs(st.session_state.lesson_content))

    combined_output = f"""📘 Lesson Plan Output:\n\n{formatted_lesson}\n\n✨ AI Customization Output:\n\n{formatted_ai}"""
    combined_output_for_docs = f"""📘 Lesson Plan Output:\n\n{formatted_lesson_for_docs}\n\n✨ AI Customization Output:\n\n{formatted_ai_for_docs}"""

    col1, col2, col3 = st.columns([1, 1, 0.5])

    with col1:
        if st.button("✏️ Edit", use_container_width=True, key="edit_btn"):
            st.session_state.edit_mode = not st.session_state.edit_mode
            st.rerun()

    with col2:
        pass

    with col3:
        download_format = st.radio(
            "📁", 
            ["DOCX", "PDF"], 
            horizontal=True, 
            label_visibility="collapsed"
        )

    with col2:
        if download_format == "DOCX":
            doc = generate_docx_file(combined_output_for_docs, title="CPALMS Lesson Plan")
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            st.download_button(
                label="⬇️ Download",
                data=doc_io,
                file_name=f"cpalms_combined_{resource_id}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        elif download_format == "PDF":
            pdf_buffer = generate_structured_pdf(combined_output)
            st.download_button(
                label="⬇️ Download",
                data=pdf_buffer,
                file_name=f"cpalms_combined_{resource_id}.pdf",
                mime="application/pdf",
                use_container_width=True
            )

if st.session_state.lesson_content:
    if st.session_state.copy_success:
        st.markdown('<div class="success-msg">✅ Content ready to copy below!</div>', unsafe_allow_html=True)
        st.session_state.copy_success = False
 
    if st.session_state.edit_mode:
        st.markdown("### ✏️ Edit Mode")

        col1, col2 = st.columns(2)

        with col1:
            edited_lesson = st.text_area(
                "Edit Lesson Plan Output:",
                value=convert_markdown_to_clean_text(st.session_state.lesson_plan_output),
                height=400,
                key="edit_lesson_plan"
            )

        with col2:
            edited_ai = st.text_area(
                "Edit AI Customization Output:",
                value=convert_markdown_to_clean_text(st.session_state.lesson_content),
                height=400,
                key="edit_ai_customization"
            )

        if st.button("💾 Save Changes", use_container_width=True):
            formatted_lesson = edited_lesson
            st.session_state.lesson_content = edited_ai
            st.session_state.edit_mode = False
            st.rerun()

    
    else:
        ai_content = st.session_state.lesson_content
        if "📘 **Previous Response**" in ai_content:
            split_parts = ai_content.split("📘 **Previous Response**")
            new_content_html = convert_markdown_to_bold_html_1(split_parts[0].strip())
            previous_content_html = convert_markdown_to_bold_html_1(split_parts[1].strip()) if len(split_parts) > 1 else ""

            st.markdown(f"""
            <style>
                .split-container {{
                    display: flex;
                    flex-direction: row;
                    gap: 20px;
                    margin-top: 30px;
                }}
                .box {{
                    flex: 1;
                    background: white;
                    border-radius: 12px;
                    padding: 25px;
                    box-shadow: 0 5px 20px rgba(0,0,0,0.08);
                    border: 2px solid #e6e6e6;
                    overflow-y: auto;
                    max-height: 700px;
                }}
                .left-label {{
                    font-weight: bold;
                    font-size: 18px;
                    margin-bottom: 10px;
                    color: #2c3e50;
                }}
                .right-label {{
                    font-weight: bold;
                    font-size: 18px;
                    margin-bottom: 10px;
                    color: #764ba2;
                }}
                .inner-box-container {{
                    display: flex;
                    flex-direction: column;
                    gap: 20px;
                }}
                .new-content-box {{
                    background: #e8f4fd;
                    padding: 18px;
                    border-radius: 10px;
                    box-shadow: 0 2px 10px rgba(52, 152, 219, 0.1);
                }}
                .previous-content-box {{
                    background: #fffce0;
                    padding: 18px;
                    border-radius: 10px;
                    box-shadow: 0 2px 10px rgba(231, 76, 60, 0.1);
                }}
            </style>
            <div class="split-container">
                <div class="box">
                    <div class="left-label">📘 Lesson Plan</div>
                    <div>{convert_markdown_to_bold_html(st.session_state.lesson_plan_output)}</div>
                </div>
                <div class="box">
                    <div class="right-label">✨ AI Customization</div>
                    <div class="inner-box-container">
                        <div class="new-content-box">
                            <div>{new_content_html}</div>
                        </div>
                        <div class="previous-content-box">
                            <strong>📘 Previous Response:</strong>
                            <div>{previous_content_html}</div>
                        </div>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)

        else:
            st.markdown("""
            <style>
                .split-container {{
                    display: flex;
                    flex-direction: row;
                    gap: 20px;
                    margin-top: 30px;
                }}
                .box {{
                    flex: 1;
                    background: white;
                    border-radius: 12px;
                    padding: 25px;
                    box-shadow: 0 5px 20px rgba(0,0,0,0.08);
                    border: 2px solid #e6e6e6;
                    overflow-y: auto;
                    max-height: 700px;
                }}
                .left-label {{
                    font-weight: bold;
                    font-size: 18px;
                    margin-bottom: 10px;
                    color: #2c3e50;
                }}
                .right-label {{
                    font-weight: bold;
                    font-size: 18px;
                    margin-bottom: 10px;
                    color: #764ba2;
                }}
            </style>
            <div class="split-container">
                <div class="box">
                    <div class="left-label">📘 Lesson Plan</div>
                    <div>{}</div>
                </div>
                <div class="box">
                    <div class="right-label">✨ AI Customization</div>
                    <div>{}</div>
                </div>
            </div>
            """.format(
                convert_markdown_to_bold_html(st.session_state.lesson_plan_output),
                convert_markdown_to_bold_html_1(st.session_state.lesson_content)
            ), unsafe_allow_html=True)



if st.session_state.lesson_content:
    st.markdown("---")
    show_history()


end_time=time.time()
total_seconds=end_time-start_time

log_query_to_blob(
            container_name="datastorage",
            resource_id=resource_id,
            benchmark_code=benchmark_code_input,
            benchmark_id=benchmark_id_input,
            query=query,
            processing_time=total_seconds,
            lesson_plan=st.session_state.lesson_plan_output,
            ai_output=st.session_state.lesson_content
        )